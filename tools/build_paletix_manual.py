from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "Manual_usuari_PaletixDesktop.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2EC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{name}"))
        if border is None:
            border = OxmlElement(f"w:{name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_indent(table, indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Title", 24, DARK_BLUE, 0, 10),
        ("Subtitle", 12, "555555", 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = style_name.startswith("Heading")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Manual d'usuari PaletixDesktop")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PaletixDesktop")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual d'usuari")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(BLUE)

    doc.add_paragraph(
        "Guia operativa per a l'aplicacio d'escriptori PaletixDesktop: navegacio, "
        "catalegs, magatzem, sincronitzacio offline, validacions i funcionalitats previstes.",
        style="Subtitle",
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    format_table(table, [2300, 7060], header=False)
    rows = [
        ("Aplicacio", "PaletixDesktop / Paletix Desktop"),
        ("Projecte", "MagatzapiV2"),
        ("Versio del manual", "1.0"),
        ("Estat de l'aplicacio", "En desenvolupament: algunes seccions son funcionals i altres son base o placeholder."),
        ("Data", "10/05/2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row.cells[0], LIGHT_BLUE)

    doc.add_paragraph()
    add_callout(
        doc,
        "Nota important",
        "Aquest manual documenta el que es pot deduir del codi actual. Com que l'aplicacio encara no esta acabada, "
        "s'han deixat apartats preparats per completar en futurs increments i s'indica quan una funcionalitat apareix "
        "com a prevista, simulada o pendent d'implementar.",
    )
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    format_table(table, [9360], header=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(f": {text}")


def format_table(table, widths: list[int], header: bool = True) -> None:
    table.autofit = False
    set_table_width(table)
    set_table_indent(table)
    set_table_borders(table)
    set_cell_margins(table)
    for row_idx, row in enumerate(table.rows):
        if header and row_idx == 0:
            set_repeat_table_header(row)
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if idx < len(widths):
                set_cell_width(cell, widths[idx])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)
            if header and row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    format_table(table, widths)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def build_manual() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    doc.add_heading("1. Objectiu i abast", level=1)
    doc.add_paragraph(
        "Aquest manual explica com utilitzar PaletixDesktop, una aplicacio d'escriptori per gestionar "
        "operacions de magatzem, cataleg comercial, clients, proveidors, stock, lots, ubicacions, flota, "
        "facturacio, gamificacio i administracio. El document esta pensat per a usuaris finals, responsables "
        "de magatzem, administracio i personal tecnic que dona suport a l'eina."
    )
    add_callout(
        doc,
        "Estat actual",
        "PaletixDesktop es troba en desenvolupament. Les vistes de productes, clients, proveidors, stock, "
        "ubicacions i lots tenen funcionalitat operativa visible. Altres moduls, com finances, flota, "
        "gamificacio, administracio i algunes parts d'operacions, disposen de base de navegacio i estructura "
        "per completar.",
    )
    doc.add_heading("1.1 Perfils destinataris", level=2)
    add_table(
        doc,
        ["Perfil", "Necessitat principal", "Parts del manual recomanades"],
        [
            ["Responsable de magatzem", "Controlar stock, ubicacions, lots, operacions pendents i sincronitzacio.", "Capitols 3, 5, 6, 9 i 11."],
            ["Operari o preparador", "Consultar existencies, registrar moviments i treballar amb mode offline si cal.", "Capitols 4, 5, 6 i 10."],
            ["Equip comercial", "Gestionar productes, clients i proveidors.", "Capitols 7 i 8."],
            ["Administracio", "Consultar facturacio, pagaments, usuaris i configuracio quan estigui complet.", "Capitols 8, 12 i 13."],
            ["Suport tecnic", "Entendre configuracio, API, base local i cua de sincronitzacio.", "Capitols 2, 10, 11 i 14."],
        ],
        [2100, 3600, 3660],
    )

    doc.add_heading("2. Requisits i posada en marxa", level=1)
    add_bullets(
        doc,
        [
            "Sistema operatiu: Windows compatible amb WinUI i l'entorn de l'aplicacio.",
            "Connexio amb l'API de Paletix: per defecte apunta a https://localhost:7137/.",
            "Clau d'API: el projecte inclou una clau per defecte i esta previst poder-la configurar.",
            "Base de dades local: PaletixDesktop utilitza SQLite amb el fitxer local paletix-local.db per cache i cua offline.",
            "Resolucio recomanada: la finestra inicial esta pensada per a 1500 x 900 i minim 1180 x 720.",
        ]
    )
    doc.add_heading("2.1 Inici de sessio actual", level=2)
    doc.add_paragraph(
        "En l'estat actual, l'aplicacio carrega una sessio local si no troba cap usuari guardat. "
        "L'usuari local per defecte es mostra com a Responsable Logistica, rol Administrador, carrec Cap de magatzem, "
        "nivell 8 i 1840 punts. Aquesta sessio serveix per provar l'aplicacio mentre el sistema d'autenticacio definitiu no esta complet."
    )
    doc.add_heading("2.2 Primer arrencada", level=2)
    add_steps(
        doc,
        [
            "Obre PaletixDesktop des de l'entorn de desenvolupament o des de l'executable publicat.",
            "Comprova que a la part superior apareix Paletix Desktop i la linia de rol de l'usuari.",
            "Mira l'indicador Online/Offline per saber si l'aplicacio esta connectada a l'API.",
            "Si hi ha canvis pendents, obre el boto de sincronitzacio pendent per revisar-los.",
            "Accedeix a Inici o a una categoria del menu superior per començar a treballar.",
        ]
    )

    doc.add_heading("3. Interficie general", level=1)
    doc.add_paragraph(
        "La finestra principal esta dividida en capcalera, menu superior de categories, barra de comandes, "
        "panell lateral de seccions i area central de treball."
    )
    add_table(
        doc,
        ["Zona", "Que mostra", "Com s'utilitza"],
        [
            ["Capcalera", "Nom de l'aplicacio, rol, botons de notificacions i canvis pendents, estat Online/Offline i punts de l'usuari.", "Serveix per controlar l'estat global abans de modificar dades."],
            ["Categories superiors", "Inici, Operacions, Magatzem, Comercial, Flota, Finances, Gamificacio i Admin.", "Selecciona el gran ambit de treball."],
            ["Panell lateral", "Seccions concretes del modul actiu, com Stock, Ubicacions o Lots.", "Canvia de vista dins del modul."],
            ["Barra de comandes", "Accions contextuals: crear, editar, eliminar, sincronitzar, taula, grid, etc.", "Executa accions sobre la vista o els registres seleccionats."],
            ["Area central", "Taules, grids, formularis laterals, dissenyadors i missatges d'estat.", "Zona principal on es consulten i editen les dades."],
        ],
        [1800, 3800, 3760],
    )
    doc.add_heading("3.1 Categories i seccions", level=2)
    add_table(
        doc,
        ["Categoria", "Seccions disponibles", "Estat funcional"],
        [
            ["Inici", "Resum executiu.", "Base preparada per al dashboard."],
            ["Operacions", "Comandes, Preparacio, Rutes.", "Navegacio i accions previstes; funcionalitat final pendent."],
            ["Magatzem", "Stock, Ubicacions, Lots.", "Funcionalitat principal implementada: consulta, alta, edicio, eliminacio, modes de vista i operacions de stock."],
            ["Comercial", "Productes, Clients, Proveidors.", "Funcionalitat de cataleg implementada per consultar i mantenir registres."],
            ["Flota", "Vehicles, Xofers.", "Base preparada per assignacions, manteniment i disponibilitat."],
            ["Finances", "Factures, Pagaments.", "Base preparada per emissio, pagaments i conciliacio."],
            ["Gamificacio", "Reptes, Premis.", "Base preparada per punts, medalles, reptes i bescanvis."],
            ["Admin", "Usuaris i rols, API i sincronitzacio.", "Base preparada per permisos, API, SQLite i diagnosi."],
        ],
        [1500, 3300, 4560],
    )

    doc.add_heading("4. Patrons comuns de treball", level=1)
    doc.add_heading("4.1 Cercar i filtrar", level=2)
    doc.add_paragraph(
        "Les vistes de cataleg i magatzem disposen d'un camp de cerca a la part superior dreta. "
        "El filtre s'aplica mentre s'escriu i actualitza el comptador de registres filtrats."
    )
    add_bullets(
        doc,
        [
            "Productes: cerca per ID, referencia, nom, descripcio o estat de sincronitzacio.",
            "Clients: cerca per ID, empresa, NIF, telefon, correu, poblacio, responsable o sincronitzacio.",
            "Proveidors: cerca per ID, empresa, marca, telefon, email, web, tipus o sincronitzacio.",
            "Stock: cerca per ID, producte, ubicacio, lot o sincronitzacio.",
            "Ubicacions: cerca per codi generat, zona, passadis, bloc, fila, columna o estat.",
            "Lots: cerca per ID, proveidor, producte, caducitat o sincronitzacio.",
        ]
    )
    doc.add_heading("4.2 Modes taula, grid i dissenyador", level=2)
    add_table(
        doc,
        ["Mode", "Quan convé", "Disponibilitat"],
        [
            ["Taula", "Quan cal comparar molts camps, ordenar visualment o revisar dades completes.", "Productes, clients, proveidors, stock, ubicacions i lots."],
            ["Grid", "Quan cal una vista mes visual i rapida de targetes.", "Productes, clients, proveidors, stock, ubicacions i lots."],
            ["Dissenyador", "Quan cal veure o crear ubicacions com una estructura visual de zona, passadis, bloc, fila i columna.", "Ubicacions; tambe s'utilitza com a selector de mapa en stock."],
        ],
        [1600, 4800, 2960],
    )
    doc.add_heading("4.3 Seleccio i edicio multiple", level=2)
    add_steps(
        doc,
        [
            "Selecciona un o mes registres des de la taula o les targetes.",
            "Comprova el text de seleccion, per exemple '2 seleccionat(s)'.",
            "Prem Editar a la barra de comandes o al control disponible.",
            "Si hi ha mes d'un registre seleccionat, utilitza les fletxes del panell lateral per passar d'un registre a l'altre.",
            "Guarda quan hagis acabat. L'aplicacio intentara sincronitzar amb l'API o deixara els canvis a la cua offline.",
        ]
    )
    doc.add_heading("4.4 Crear, guardar, cancelar i eliminar", level=2)
    add_bullets(
        doc,
        [
            "Crear obre un panell lateral amb els camps buits o valors inicials.",
            "Guardar valida els camps obligatoris i els formats abans d'aplicar el canvi.",
            "Cancelar tanca el panell i mostra un missatge d'edicio cancelada.",
            "Eliminar requereix que hi hagi registres seleccionats; si no n'hi ha, l'aplicacio mostra un missatge d'avís.",
            "Les operacions fetes sense connexio poden quedar marcades com a pendents.",
        ]
    )

    doc.add_heading("5. Modul Magatzem", level=1)
    doc.add_heading("5.1 Stock", level=2)
    doc.add_paragraph(
        "La vista Stock mostra existencies operatives i permet registrar entrades, moviments, ajustos i reserves. "
        "Els indicadors principals son Total, Filtrats, Reservats i Baix disponible."
    )
    add_table(
        doc,
        ["Camp visible", "Descripcio"],
        [
            ["ID", "Identificador del registre de stock."],
            ["Producte", "Producte associat a l'existencia."],
            ["Ubicacio", "Ubicacio fisica del magatzem."],
            ["Lot", "Lot associat, si n'hi ha."],
            ["Total", "Unitats totals en stock."],
            ["Reservat", "Unitats reservades per comandes."],
            ["Disponible", "Unitats disponibles calculades."],
            ["Sincronitzacio", "Estat local, pendent, error o sincronitzat."],
        ],
        [2200, 7160],
    )
    doc.add_heading("5.1.1 Crear o editar stock", level=3)
    add_steps(
        doc,
        [
            "Entra a Magatzem > Stock.",
            "Prem Entrada o crea un registre segons l'accio disponible.",
            "Tria Producte, Ubicacio i, opcionalment, Lot.",
            "Indica Total en stock i Reservat per comandes.",
            "Guarda. Si el total reservat supera el total en stock, l'aplicacio mostra error de validacio.",
        ]
    )
    doc.add_heading("5.1.2 Operacions de stock", level=3)
    add_table(
        doc,
        ["Operacio", "Que fa", "Dades necessaries"],
        [
            ["Entrada de stock", "Registra entrada fisica de producte en una ubicacio.", "Producte, ubicacio, lot opcional, quantitat i motiu opcional."],
            ["Moure stock", "Mou unitats d'una ubicacio origen a una ubicacio desti.", "Producte, origen, desti, lot opcional, quantitat i motiu."],
            ["Ajust inventari", "Canvia el total d'un registre seleccionat.", "Stock seleccionat, nou total i motiu."],
            ["Reservar stock", "Incrementa la quantitat reservada per comandes.", "Stock seleccionat, quantitat i motiu."],
            ["Alliberar reserva", "Redueix reserva i torna unitats a disponible.", "Stock seleccionat, quantitat i motiu."],
            ["Historial", "Consulta moviments de stock registrats.", "Cap, excepte tenir dades carregades."],
        ],
        [1800, 3700, 3860],
    )
    doc.add_heading("5.1.3 Selector de mapa", level=3)
    doc.add_paragraph(
        "En formularis de stock apareix el boto Mapa per seleccionar ubicacions visualment. "
        "Aquest selector aprofita el dissenyador d'ubicacions i ajuda a evitar errors en codis de zona, passadis, bloc, fila i columna."
    )

    doc.add_heading("5.2 Ubicacions", level=2)
    doc.add_paragraph(
        "La vista Ubicacions representa l'estructura fisica del magatzem. Una ubicacio es defineix per zona, passadis, bloc d'estanteria, fila i columna."
    )
    add_table(
        doc,
        ["Indicador o camp", "Descripcio"],
        [
            ["Total", "Nombre total d'ubicacions carregades."],
            ["Filtrats", "Ubicacions que coincideixen amb la cerca."],
            ["Zones", "Nombre de zones diferents."],
            ["Disponibles", "Registres no marcats com a eliminacio pendent."],
            ["Codi", "Codi generat a partir de coordenades, per exemple Z1-P2-B1-F3-C4."],
        ],
        [2200, 7160],
    )
    doc.add_heading("5.2.1 Alta manual", level=3)
    add_steps(
        doc,
        [
            "Obre Magatzem > Ubicacions.",
            "Prem Nova ubicacio.",
            "Introdueix Zona, Passadis, Bloc estanteria, Fila i Columna.",
            "Guarda el registre.",
        ]
    )
    doc.add_heading("5.2.2 Generador d'ubicacions", level=3)
    add_paragraph_text = (
        "El generador crea moltes ubicacions a partir de rangs. Permet indicar una zona i intervals de passadissos, blocs, files i columnes. "
        "L'aplicacio calcula una vista previa i limita l'operacio a 1000 ubicacions per execucio. Les ubicacions ja existents se salten."
    )
    doc.add_paragraph(add_paragraph_text)
    add_steps(
        doc,
        [
            "Prem Generar dins d'Ubicacions.",
            "Indica la zona i els rangs inicial/final de passadis, bloc, fila i columna.",
            "Revisa el resum i els errors de validacio, si n'hi ha.",
            "Prem Generar. Durant el proces es mostra el progres.",
            "Revisa el dissenyador per confirmar que la nova estructura apareix correctament.",
        ]
    )
    add_callout(
        doc,
        "Criteri de seguretat",
        "No generis rangs amplis sense revisar la vista previa. Si el magatzem real encara no esta modelat, acorda primer la nomenclatura de zones i passadissos.",
    )

    doc.add_heading("5.3 Lots de proveidor", level=2)
    doc.add_paragraph(
        "La vista Lots registra entrades de proveidor, producte, quantitat rebuda i dates de demanat, recepcio i caducitat."
    )
    add_table(
        doc,
        ["Camp", "Us"],
        [
            ["Proveidor", "Origen del lot. Es tria del desplegable de proveidors."],
            ["Producte", "Producte rebut. Es tria del desplegable de productes."],
            ["Quantitat rebuda", "Unitats rebudes; ha de ser igual o superior a 1."],
            ["Data demanat", "Opcional; data en que es va demanar el lot."],
            ["Data rebut", "Data de recepcio."],
            ["Data caducitat", "No pot ser anterior a la data de recepcio."],
        ],
        [2200, 7160],
    )

    doc.add_heading("6. Modul Comercial", level=1)
    doc.add_heading("6.1 Productes", level=2)
    doc.add_paragraph(
        "Productes es el cataleg principal. Permet consultar, crear, editar, eliminar i importar registres, "
        "amb vista de taula o grid i imatge opcional."
    )
    add_table(
        doc,
        ["Camp", "Obligatori", "Validacio o observacio"],
        [
            ["Referencia", "No", "Maxim 50 caracters; nomes lletres, numeros, punts, guions, barres i guio baix, començant per lletra o numero."],
            ["Nom", "Si", "Maxim 150 caracters."],
            ["Descripcio", "No", "Maxim 300 caracters."],
            ["Tipus de producte", "Si", "S'ha de triar del desplegable."],
            ["Volum (ml)", "No", "Numero igual o superior a 0."],
            ["Proveidor", "Si", "S'ha de triar del desplegable."],
            ["Ubicacio", "Si", "S'ha de triar del desplegable."],
            ["Caixes per palet", "Si", "Enter igual o superior a 1."],
            ["URL imatge", "No", "URL http/https valida, sense espais."],
            ["Actiu", "Si", "Actiu o Inactiu."],
            ["Preu venda caixa", "Si", "Numero igual o superior a 0."],
            ["Cost per caixa", "Si", "Numero igual o superior a 0."],
            ["Estabilitat al palet", "No", "Enter igual o superior a 0."],
            ["Pes (kg)", "No", "Numero igual o superior a 0."],
        ],
        [2200, 1200, 5960],
    )
    doc.add_heading("6.1.1 Importacio manual", level=3)
    doc.add_paragraph(
        "L'accio Importar crea un producte d'exemple amb referencia IMP i dades inicials. "
        "Aquesta funcio sembla preparada com a increment temporal per simular una importacio fins que existeixi un importador complet de fitxers."
    )

    doc.add_heading("6.2 Clients", level=2)
    doc.add_paragraph(
        "Clients gestiona fitxes comercials amb dades de contacte, adreca, poblacio i responsable. "
        "La vista mostra comptadors de Total, Filtrats, Amb email i Pendents."
    )
    add_table(
        doc,
        ["Camp", "Obligatori", "Validacio o observacio"],
        [
            ["Empresa", "Si", "Nom de l'empresa client."],
            ["NIF empresa", "No", "Format DNI, NIE o CIF, per exemple 12345678Z, X1234567L o B12345678."],
            ["Telefon", "Si", "6 a 15 digits; admet +, espais, punts, guions i parentesis."],
            ["Correu electronic", "No", "Format de correu valid."],
            ["Adreca", "Si", "Maxim 500 caracters."],
            ["Poblacio", "Si", "Maxim 100 caracters."],
            ["Responsable", "No", "Maxim 255 caracters."],
        ],
        [2200, 1200, 5960],
    )
    doc.add_heading("6.2.1 Historial de client", level=3)
    doc.add_paragraph(
        "L'accio Historial esta prevista. Actualment, si no hi ha cap client seleccionat, demana seleccionar-ne un; "
        "si n'hi ha, informa que l'historial esta pendent d'implementar."
    )

    doc.add_heading("6.3 Proveidors", level=2)
    doc.add_paragraph(
        "Proveidors manté el directori de proveidors i dades comercials relacionades amb tipus de producte principal."
    )
    add_table(
        doc,
        ["Camp", "Obligatori", "Validacio o observacio"],
        [
            ["Marca matriu", "No", "Maxim 100 caracters."],
            ["Empresa", "Si", "Maxim 100 caracters."],
            ["Telefon", "Si", "Maxim 16 caracters; format telefonic valid."],
            ["Correu electronic", "Si", "Maxim 200 caracters; format valid."],
            ["Adreca", "No", "Adreca fiscal o magatzem."],
            ["Enllac web", "No", "URL http/https valida, sense espais."],
            ["Tipus producte principal", "No", "Si s'indica, ha d'existir al desplegable."],
        ],
        [2200, 1200, 5960],
    )
    doc.add_heading("6.3.1 Comparar proveidors", level=3)
    doc.add_paragraph(
        "L'accio Comparar esta prevista. Actualment demana seleccionar dos o mes proveidors i mostra un missatge indicant que la comparacio esta pendent."
    )

    doc.add_heading("7. Operacions", level=1)
    doc.add_paragraph(
        "El modul Operacions esta estructurat per gestionar comandes, preparacio i rutes. "
        "Les accions visibles son Nova, Assignar, Validar, Iniciar, Pausar, Incidencia, Planificar i Optimitzar. "
        "Aquestes pantalles encara s'han de completar, pero la navegacio ja reserva l'espai funcional."
    )
    add_table(
        doc,
        ["Seccio", "Objectiu previst", "Accions previstes"],
        [
            ["Comandes", "Creacio, estat i seguiment de comandes.", "Nova, Assignar, Validar."],
            ["Preparacio", "Picking, verificacio i incidencies.", "Iniciar, Pausar, Incidencia."],
            ["Rutes", "Assignacio de vehicle, xofer i entregues.", "Planificar, Optimitzar."],
        ],
        [1800, 4300, 3260],
    )

    doc.add_heading("8. Flota, finances i gamificacio", level=1)
    add_table(
        doc,
        ["Modul", "Seccions", "Funcionalitat prevista"],
        [
            ["Flota", "Vehicles, Xofers.", "Disponibilitat, tipus, manteniment, assignacions, torns i rendiment."],
            ["Finances", "Factures, Pagaments.", "Emissio, estat, venciments, cobraments, conciliacio i exportacio."],
            ["Gamificacio", "Reptes, Premis.", "Objectius, punts, participants, medalles, bescanvis i recompenses."],
        ],
        [1600, 2500, 5260],
    )
    doc.add_paragraph(
        "Aquestes seccions mostren base de pagina i metriques simulades o preparades. Quan s'implementin les pantalles finals, "
        "caldra afegir procediments detallats d'alta, edicio, validacio, tancament i informes."
    )

    doc.add_heading("9. Inici i dashboard", level=1)
    doc.add_paragraph(
        "La categoria Inici conté el Resum executiu. Actualment esta preparat com a espai principal per construir el dashboard. "
        "El cataleg de moduls inclou metriques i activitats d'exemple per operacions, magatzem, productes, clients, proveidors, "
        "flota, facturacio, gamificacio, usuaris i administracio."
    )
    add_table(
        doc,
        ["Tipus d'informacio", "Exemples actuals", "Us previst"],
        [
            ["Metriques", "Pendents, stock disponible, clients actius, factures obertes, reptes actius.", "Donar visio rapida de l'estat operatiu."],
            ["Activitats", "Comandes, lots pendents, revisions de permisos, incidencies.", "Prioritzar tasques del dia."],
            ["Accions principals", "Nova comanda, ajustar stock, nou producte, sincronitzar.", "Arribar rapidament a les accions habituals."],
        ],
        [1900, 3900, 3560],
    )

    doc.add_heading("10. Sincronitzacio i mode offline", level=1)
    doc.add_paragraph(
        "PaletixDesktop esta dissenyat per treballar amb API i cache local SQLite. Quan l'API esta disponible, carrega dades remotes i actualitza el cache. "
        "Quan no hi ha connexio o una operacio no es pot enviar, el canvi pot quedar a la cua de sincronitzacio."
    )
    doc.add_heading("10.1 Indicadors", level=2)
    add_table(
        doc,
        ["Indicador", "Significat", "Accio recomanada"],
        [
            ["Online", "L'aplicacio considera que la connexio esta disponible.", "Treballar normalment i revisar pendents si n'hi ha."],
            ["Offline", "No s'ha pogut carregar o sincronitzar amb l'API.", "Continuar si cal amb dades locals i revisar la cua quan torni la connexio."],
            ["Canvis pendents", "Hi ha operacions locals esperant sincronitzar o amb error.", "Obrir el desplegable, reintentar o descartar segons el cas."],
            ["Estat per registre", "Cada registre pot mostrar sincronitzat, pendent, error o eliminacio pendent.", "Evitar duplicar accions sobre registres amb error fins revisar-los."],
        ],
        [1900, 3900, 3560],
    )
    doc.add_heading("10.2 Revisar canvis pendents", level=2)
    add_steps(
        doc,
        [
            "Prem el boto de canvis pendents a la capcalera.",
            "Llegeix l'estat general i la llista d'operacions.",
            "Per un canvi concret, prem Reintentar si vols tornar-lo a enviar.",
            "Prem Descartar nomes si confirmes que no cal conservar el canvi local.",
            "Utilitza Reintentar tot quan la connexio s'hagi recuperat i vulguis reenviar tots els errors.",
        ]
    )
    doc.add_heading("10.3 Entitats sincronitzades", level=2)
    add_bullets(
        doc,
        [
            "Clients.",
            "Proveidors.",
            "Productes.",
            "Stock.",
            "Operacions de stock.",
            "Ubicacions.",
            "Lots de proveidor.",
        ]
    )

    doc.add_heading("11. Permisos i visibilitat", level=1)
    doc.add_paragraph(
        "La navegacio es filtra segons rol i carrec. L'administrador pot accedir a tot. Altres perfils veuen nomes els moduls que encaixen amb la seva responsabilitat."
    )
    add_table(
        doc,
        ["Perfil detectat", "Acces principal"],
        [
            ["Administrador o admin", "Totes les funcionalitats."],
            ["Cap, responsable o gestor", "Operacions, magatzem, cataleg, clients, proveidors, flota i gamificacio."],
            ["Preparador, magatzem o operari", "Operacions, magatzem i gamificacio."],
            ["Xofer o transport", "Operacions, flota i gamificacio."],
            ["Finances o administracio", "Facturacio, clients i proveidors."],
            ["Sense usuari o perfil no reconegut", "Sense acces excepte dashboard quan correspongui."],
        ],
        [2700, 6660],
    )

    doc.add_heading("12. Administracio i configuracio", level=1)
    doc.add_paragraph(
        "El modul Admin agrupa Usuaris i rols i API i sincronitzacio. Actualment funciona com a base de pantalla i cataleg d'accions. "
        "Les metriques internes indiquen API local, SQLite actiu i permisos aplicats a la navegacio."
    )
    add_table(
        doc,
        ["Element", "Valor o estat actual"],
        [
            ["URL API per defecte", "https://localhost:7137/"],
            ["Base local", "paletix-local.db"],
            ["Interval de comprovacio de connexio", "8 segons."],
            ["Timeout de comprovacio", "3 segons."],
            ["Timeout de peticio", "30 segons."],
            ["Finestra inicial", "1500 x 900; minim 1180 x 720."],
        ],
        [3000, 6360],
    )
    add_callout(
        doc,
        "Nota tecnica",
        "El codi actual intenta llegir variables d'entorn utilitzant com a nom el valor de la URL i de la clau, no noms com PALETIX_API_URL o PALETIX_API_KEY. "
        "Al manual d'administracio tecnica futur caldria documentar el mecanisme definitiu quan estigui corregit o estabilitzat.",
    )

    doc.add_heading("13. Missatges, validacions i errors", level=1)
    doc.add_heading("13.1 Validacions habituals", level=2)
    add_table(
        doc,
        ["Situacio", "Missatge o comportament", "Solucio"],
        [
            ["Camp obligatori buit", "El formulari mostra que el camp es obligatori.", "Omple el camp indicat abans de guardar."],
            ["Telefon invalid", "S'indica que el telefon ha de tenir format valid.", "Utilitza 6 a 15 digits i nomes simbols permesos."],
            ["Email invalid", "S'indica que el correu electronic ha de tenir format valid.", "Revisa domini, arrova i espais."],
            ["URL invalid", "S'indica que la URL ha de ser http/https i sense espais.", "Escriu una URL completa com https://exemple.cat."],
            ["Lookup invalid", "S'indica que s'ha de triar del desplegable.", "Selecciona un valor existent."],
            ["Quantitat negativa o zero", "S'indica el minim permes.", "Introdueix una quantitat igual o superior al minim."],
            ["Caducitat anterior a recepcio", "No es permet guardar el lot.", "Canvia data de caducitat o data de recepcio."],
        ],
        [2300, 3600, 3460],
    )
    doc.add_heading("13.2 Bones practiques", level=2)
    add_bullets(
        doc,
        [
            "Abans d'eliminar, comprova que la seleccio sigui correcta.",
            "Despres d'una sessio offline, revisa sempre els canvis pendents.",
            "Utilitza el mode taula per auditories i el mode grid per revisio rapida.",
            "En ubicacions, acorda una convencio de zones i passadissos abans d'utilitzar el generador.",
            "En stock, escriu un motiu clar als ajustos i moviments quan el camp estigui disponible.",
        ]
    )

    doc.add_heading("14. Funcionalitats pendents i estructura futura", level=1)
    add_table(
        doc,
        ["Area", "Pendent de completar", "Proposta per ampliar el manual"],
        [
            ["Autenticacio", "Login real, canvi d'usuari, recuperacio de sessio.", "Afegir capitol d'acces, perfils i tancament de sessio."],
            ["Operacions", "Comandes, picking, incidencies i rutes completes.", "Afegir procediments de cicle de comanda."],
            ["Flota", "Vehicles, xofers, manteniments i assignacions reals.", "Afegir gestio de disponibilitat i manteniment."],
            ["Finances", "Factures, pagaments, venciments i exportacions.", "Afegir circuit de facturacio i conciliacio."],
            ["Gamificacio", "Reptes, medalles, premis i bescanvis operatius.", "Afegir regles de punts, aprovacions i recompenses."],
            ["Administracio", "Pantalles reals d'usuaris, rols, API, logs i diagnosi.", "Afegir manual d'administrador."],
            ["Importacions", "Importacio massiva real de productes o altres dades.", "Afegir formats acceptats, errors i validacio de fitxers."],
            ["Informes", "Exportacions, informes i dashboards.", "Afegir lectura de KPI i generacio de documents."],
        ],
        [1700, 3900, 3760],
    )

    doc.add_heading("15. Annex A: resum rapid per modul", level=1)
    add_table(
        doc,
        ["Modul", "Consultar", "Crear/Editar", "Eliminar", "Notes"],
        [
            ["Productes", "Si", "Si", "Si", "Inclou importacio manual temporal i imatge URL."],
            ["Clients", "Si", "Si", "Si", "Historial pendent."],
            ["Proveidors", "Si", "Si", "Si", "Comparacio pendent."],
            ["Stock", "Si", "Si", "Si", "Inclou entrada, moviment, ajust, reserva, alliberament i historial."],
            ["Ubicacions", "Si", "Si", "Si", "Inclou dissenyador i generador massiu."],
            ["Lots", "Si", "Si", "Si", "Control de caducitat i quantitat rebuda."],
            ["Operacions", "Base", "Pendent", "Pendent", "Navegacio i accions definides."],
            ["Flota", "Base", "Pendent", "Pendent", "Pantalles finals pendents."],
            ["Finances", "Base", "Pendent", "Pendent", "Pantalles finals pendents."],
            ["Gamificacio", "Base", "Pendent", "Pendent", "Pantalles finals pendents."],
            ["Admin", "Base", "Pendent", "Pendent", "Pantalles finals pendents."],
        ],
        [1700, 1400, 1500, 1300, 3460],
    )

    doc.add_heading("16. Annex B: glossari", level=1)
    add_table(
        doc,
        ["Terme", "Definicio"],
        [
            ["API", "Servei amb el qual PaletixDesktop sincronitza dades."],
            ["Cache local", "Copia local de dades guardada en SQLite per accelerar carrega i permetre treball offline."],
            ["Cua offline", "Llista d'operacions locals pendents d'enviar a l'API."],
            ["Grid", "Vista de targetes per revisar registres visualment."],
            ["Taula", "Vista de files i columnes per comparar camps."],
            ["Dissenyador", "Vista visual de les ubicacions del magatzem."],
            ["Lot", "Entrada de producte associada a proveidor, quantitat i caducitat."],
            ["Reserva", "Quantitat de stock separada per comandes."],
            ["Alliberament", "Accio que retorna stock reservat a disponible."],
            ["Sincronitzat", "Registre que no te canvis locals pendents."],
        ],
        [2300, 7060],
    )

    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build_manual()
